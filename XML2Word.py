#!/usr/bin/python
# -*- coding: utf-8 -*-

from lxml import etree as ET
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from collections import Counter
from docx.enum.text import WD_BREAK
import sys
from docx.enum.style import WD_STYLE_TYPE



#APPROVED::::Function to write a line to docx, specifying font characteristics
def writeText(p1, textToWrite1, italicText, boldText, underlineText, font_size0, inline0, font_name0):
	if(textToWrite1 == ("*** NEW PAGE HERE ***")):
		run = p1.add_run()
		run.add_break(WD_BREAK.PAGE)
		return
	if (inline0 == "False"):
		textToWrite1 = " " + textToWrite1	
	run = p1.add_run( textToWrite1)
	fontUsed = run.font
	if(italicText):
		fontUsed.italic = True
	if(boldText):
		fontUsed.bold = True
	if(underlineText):		
		fontUsed.underline = True
	fontUsed.name = font_name0
	if ((int(float(font_size0)) > 0) and (int(float(font_size0)) < 1000)):
		fontUsed.size = Pt(int(float(font_size0)))
	return


def buildString(formattings_element, charParams_element):
	
	strBuilt = ""
	bb = len(formattings_element.findall(charParams_element))
	
	if ( bb < 1):
		#print (formattings_element.findall('charParams'))
		
		return formattings_element.text
	else:
		for a_charParams in formattings_element.iter(charParams_element):
			strBuilt = strBuilt + a_charParams.text
			#print("found Xter ", a_charParams)
		return strBuilt
	
#Writes Paragraphs with the original style into docx document
def writeParagraphtoDocument(root_element, page_element, pars_element, formattings_element, fileName, charParams_element):	
	docx_document = Document(fileName)
	stringToWrite = ""
	line_spa = 111
	replaced_flag = 0
	init_page = 5555
	current_page = 5555
	first_run = 1;
	pagenum_detected = 0

	num_of_pages = root_element.get("NUMBEROFPAGES", 0)	
	pagenum_location = root_element.get("PAGENUMBERINGLOCATION", 0)
	x = 0

	#paragraph = document1.add_paragraph()
	#paragraph.style = document1.styles['BfA Normal']
	#run = paragraph.add_run( "THIS IS BFA NORMAL")


	for a_par in root_element.iter(pars_element):
		print("Writing Paragraph Number ", x)
		x += 1
		line_spa1 = a_par.get("lineSpacing", 111)
		alignment1 = a_par.get("align", "NONE")
		p = setParagraph(docx_document, line_spa1, alignment1)

		#Set Paragraph Style	
		par_style = a_par.get("STYLE")
		print("STYLE DETECTED IS ... ", par_style)		
		p.style = par_style

		if((line_spa != line_spa1)or(alignment != alignment1)):
			stringToWrite = "Current Line Spacing XXX is " + str(line_spa1) + " Current Alignment YYY is " + str(alignment1)
			line_spa = line_spa1
			alignment = alignment1
		llb_bold = 0;
		llb_italic = 0;
		llb_previous = 0;
		
		for a_format in a_par.iter(formattings_element):
			font_size1 = a_format.get("fs", 222)
			isBold1 = a_format.get("bold", 0)
			isItalic1 = a_format.get("italic", 0)
			font_name00 = a_format.get("ff", "Calibri")

			inline_1 = a_format.get("INLINE", "False")

			
			#stringToWrite = a_format.text
			stringToWrite = buildString(a_format,charParams_element)


			if(replaced_flag):
				inline_1 = "True"

			#Replace Word Multi Line Continuation Character Here
			if (stringToWrite.find("¬") > 0):
				stringToWrite = stringToWrite.replace("¬", "")
				replaced_flag = 1
			elif (stringToWrite.find("¬") < 0):
				replaced_flag = 0

			pg1 = list(a_format.iterancestors(page_element))			
			pg2 = pg1[0]
			current_page = pg2.get("PAGENUMBERINDEX", 0)

			if((current_page != init_page)and (first_run != 1)):	
				writeText(p,"*** NEW PAGE HERE ***", isItalic1, isBold1,0, font_size1, inline_1, font_name00)

				if((pagenum_location == "TOP") or (pagenum_location == "BOTTOM")):
					writeText(p,"*** CURRENT PAGE NUMBERING  DETECTED IS *** " + pg2.get("PAGENUMBERSCANNED", 0) + "\n", 0, 0,0, 10, inline_1, font_name00)
			if(((pagenum_location == "TOP") or (pagenum_location == "BOTTOM")) and (first_run == 1)):
					writeText(p,"*** CURRENT PAGE NUMBERING  DETECTED IS *** " + pg2.get("PAGENUMBERSCANNED", 0) + "\n", 0, 0,0, 10, inline_1, font_name00)			



				#init_page = current_page
			init_page = current_page # JUST ADDED

			first_run = 0
			writeText(p, stringToWrite, isItalic1, isBold1,0, font_size1, inline_1, font_name00)
	docx_document.save(fileName)


#APPROVED::::Sets Paragraph based on the original document
def setParagraph(D2, lineSpacing2, alignment2):
	p2 = D2.add_paragraph("")
	p2Format = p2.paragraph_format
	if(alignment2 == "justified"):
		p2Format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	return p2

#Strip Digits from the beginning and end of first three lines, to check if there are headers plus page/section numbering
def stripDigits(string2a):
	string2a = str.lstrip(string2a, ' 0123456789 ')
	string2a = str.rstrip(string2a, ' 0123456789 ')
	return string2a

#Strips Non Digits from String to Search for Page Numbering
def stripNONDigits_3(string1a, pgNum):
	string2a = str(string1a)
	newString = ""
	pg_1 = 0
	#num_of_pages_generic = getNumberOfPages(root, pages, formattings)

	for char_1 in string2a:
		if not(char_1.isdigit()):
			if (pg_1):
				#return newString
				break # Added			
		else:
			newString+=char_1;
			pg_1 = 1
	if((newString == "")or(int(newString) > (pgNum * 10))):
		return ""
	elif(int(newString) > 0):
		# and (int(newString) <= pgNum)):
		return newString
	

#Create a matrix of number of pages X 3 for ease of header comparison and checks
def create2DimArray(headerArray, numOfPages):
	pageLineMatrix = [[0 for i in range(3)] for i in range(numOfPages)]
	row = 0		
	ccount = 0
	for page_1 in range(numOfPages):
		col = 0
		for line_1 in range(3):
			pageLineMatrix[row][col] = headerArray[ccount]
			ccount += 1
			col += 1
		row += 1
	return pageLineMatrix

#Take a 2 dim array of page and 3 top/bottom lines and return a dictionary of common lines across pages	
def compare3TopBottomLines(pageLinesMatrix, pageCount):
	foundHeaderAlready = 0
	headerList = list()
	commonLines_dict = dict()
	row_1 = 0
	col_1 = 0
	headerCount_1 = 0 #To Check if First Line is Header
	headerCount2_1 = 0
	headerCount3_1 = 0		
	for ccount_1 in range(pageCount):
		if(pageLinesMatrix[row_1][col_1] == pageLinesMatrix[row_1 + 2][col_1]):
			headerCount_1 += 1
			commonLines_dict[ccount_1] = pageLinesMatrix[row_1][col_1]
		if(pageLinesMatrix[row_1][col_1 + 1] == pageLinesMatrix[row_1 + 2][col_1 + 1]):
			headerCount2_1 += 1
			commonLines_dict[ccount_1] = pageLinesMatrix[row_1][col_1 + 1]
		if(pageLinesMatrix[row_1][col_1 + 2] == pageLinesMatrix[row_1 + 2][col_1 + 2]):
			headerCount3_1 += 1
			commonLines_dict[ccount_1] = pageLinesMatrix[row_1][col_1 + 2]
		if(row_1 >= (pageCount - 3)): #Check to only compare 3 lines else array overflow
			break
		else:
			row_1 += 1
	return commonLines_dict

# Checker for presence of Headers
#In Progresss
def getHeader (header3, pages):
	header33 = list() #Aim is not to modify the list of headers
	header33 = header3
	if(len(header3) != (3*pages)):
		print("Header Extraction Incorrect ")
		return
	if(pages < 2):	
		print("Only Single Page ")
		return
	#Strip digits at start and end of top 3 lines	
	else:
		headerIndex = 0
		for eachLine in header33:						
			header33[headerIndex] = stripDigits(eachLine)
			print("IE", header33[headerIndex])
			headerIndex += 1
		pageMatrix = create2DimArray(header33, pages)
		dictOfHeaders = compare3TopBottomLines(pageMatrix, pages)
		print("DICT OF HEADERS IS ****    ", dictOfHeaders)
	return dictOfHeaders

#Check for the Presence of Footers
# In Progress
def getFooter (footer4, pages):	
	if(len(footer4) != (3*pages)):
		print("Footer Extraction Incorrect ")
		return
	if(pages < 2):	
		print("Only Single Page ")
		return
	#Strip digits at start and end of top 3 lines	
	else:
		footerIndex = 0
		for eachLine in footer4:
			footer4[footerIndex] = stripDigits(eachLine)
			print ("THAT WAS NEW FOOTER FUNCTION ", footer4[footerIndex])
			footerIndex += 1
		pageMatrix = create2DimArray(footer4, pages)
		dictOfFooters = compare3TopBottomLines(pageMatrix, pages)
		print("DICT OF FOOTERS IS ****    ", dictOfFooters)
	return dictOfFooters

def filterNumericStrings_3(string1a):
	string2a = str(string1a)
	if(string2a.isdigit()):
		return string2a
	else:
		return ""
		
#Extract Top 3 Lines from the XML Using Elements - formatting, 
def getTop3Line(root_element, pages_element, formattings_element):
	num1 = getNumberOfPages(root_element, pages_element, formattings_element) * 3
	top3linesandPage = [" "] * num1 
	numOfPagesCounted = 0
	for a_page_element in root_element.iter(pages_element):
		count3Lines_1 = 0		
		for a_format in a_page_element.iter(formattings_element):		
			if (count3Lines_1 > 2):
				break
			elif (count3Lines_1 <= 2):				
				top3linesandPage[numOfPagesCounted * 3 + count3Lines_1] = a_format.text
			count3Lines_1 = count3Lines_1 + 1			
		numOfPagesCounted += 1
	return top3linesandPage

#Get Bottom three lines from each Page
def getBottom3Line(root_element, pages_element, formattings_element):
	page_num = 0
	num1 = getNumberOfPages(root_element, pages_element, formattings_element) * 3
	bottom3linesandPage = [" "] * num1 
	for a_page33 in root_element.iter(pages_element):
		countPageLines = 0		
		for a_format33 in a_page33.iter(formattings_element):		
			countPageLines = countPageLines + 1
		bottom3Counter = 0
		hcounter = 0
		for a_format333 in a_page33.iter(formattings_element):
			if((bottom3Counter > countPageLines - 4)and(bottom3Counter < countPageLines + 1)):
				bottom3linesandPage[page_num * 3 + hcounter] = a_format333.text
				hcounter += 1
			bottom3Counter += 1
		page_num += 1
	return bottom3linesandPage

# Get the number of Pages on the XML file
def getNumberOfPages(root_element, pages_element, formattings_element):
	numOfPagesCounted = 0
	for a_page_element in root_element.iter(pages_element):
		numOfPagesCounted += 1
	return numOfPagesCounted

def addHeadingStyles(doc_1):
	styles = doc_1.styles
	styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH, builtin=True)
	styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH, builtin=True)
	styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH, builtin=True)
	styles.add_style('Heading 4', WD_STYLE_TYPE.PARAGRAPH, builtin=True)
	styles.add_style('Heading 5', WD_STYLE_TYPE.PARAGRAPH, builtin=True)
	styles.add_style('Heading 6', WD_STYLE_TYPE.PARAGRAPH, builtin=True)





#END OF FUNCTIONS,
def main():
	#Define XML File Parameters
	
	print("Checking I/O Files ...")
	xml_file_name = ""
	if(sys.argv[1].rfind(".xml") == -1):
		xml_file_name = sys.argv[1] + ".xml"
	else:
		xml_file_name = sys.argv[1]
	print("Input File is ...", xml_file_name)
	print("*********************************")

	tree = ET.parse(xml_file_name)
	print("PARSED XML FILE ...", tree)
	print("*********************************")
	root = tree.getroot()
	nameSpace = "{http://www.abbyy.com/FineReader_xml/FineReader10-schema-v1.xml}"
	blocks = nameSpace + "block"
	formattings = nameSpace + "formatting"
	charParams = nameSpace + "charParams"
	pars = nameSpace + "par"
	lines = nameSpace + "line" 
	pages = nameSpace + "page"
	texts = nameSpace + "text"
	


	#Define Docx Parameters and Document to write to	
	created_file = ""
	if(sys.argv[2].rfind(".docx") == -1):
		created_file = sys.argv[2] + ".docx"
	else:
		created_file = sys.argv[2]
	print("Output File is ...", created_file)
	print("*********************************")
	print("Processing ...")

	document1 = Document("Template.docx")
	addHeadingStyles(document1)
	document1.save(created_file)	
	#Write XMl Contents to docx
	print("Started Writing to Word Document")
	writeParagraphtoDocument(root, pages, pars, formattings, created_file, charParams)
main()




















